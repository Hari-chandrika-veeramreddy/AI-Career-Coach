import os
try:
    import pymupdf as fitz
except ImportError:
    import fitz
import docx

def parse_pdf(file_bytes_or_path):
    """Extract clean text from a PDF file using PyMuPDF."""
    text = ""
    try:
        if isinstance(file_bytes_or_path, (str, bytes)):
            if isinstance(file_bytes_or_path, str) and os.path.exists(file_bytes_or_path):
                doc = fitz.open(file_bytes_or_path)
            else:
                doc = fitz.open(stream=file_bytes_or_path, filetype="pdf")
            
            for page in doc:
                text += page.get_text("text") + "\n"
            doc.close()
    except Exception as e:
        print(f"Error parsing PDF: {e}")
        return ""
    return text.strip()

def parse_docx(file_bytes_or_path):
    """Extract clean text from a DOCX file using python-docx."""
    text = ""
    try:
        if isinstance(file_bytes_or_path, str) and os.path.exists(file_bytes_or_path):
            doc = docx.Document(file_bytes_or_path)
        else:
            from io import BytesIO
            doc = docx.Document(BytesIO(file_bytes_or_path))

        full_text = []
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text.strip())
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        full_text.append(cell.text.strip())
        text = "\n".join(full_text)
    except Exception as e:
        print(f"Error parsing DOCX: {e}")
        return ""
    return text.strip()

def extract_resume_text(uploaded_file):
    """
    Main entry point for extracting text from uploaded resume files.
    Supports PDF and DOCX formats. Handles empty or corrupted files safely.
    """
    if uploaded_file is None:
        return ""

    file_name = uploaded_file.name.lower()
    try:
        file_bytes = uploaded_file.read()
        if not file_bytes:
            return ""

        if file_name.endswith('.pdf'):
            return parse_pdf(file_bytes)
        elif file_name.endswith('.docx'):
            return parse_docx(file_bytes)
        elif file_name.endswith('.txt'):
            return file_bytes.decode('utf-8', errors='ignore').strip()
        else:
            return ""
    except Exception as e:
        print(f"Error extracting resume text: {e}")
        return ""
